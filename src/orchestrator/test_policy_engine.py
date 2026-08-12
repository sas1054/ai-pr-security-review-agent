import pytest
import httpx
from io import BytesIO
from docx import Document
from pypdf import PdfWriter

from policy_engine import (
    AzureOpenAIPolicyInterpreter,
    PolicyEngineError,
    _validate_public_https,
    compile_proposal,
    extract_document,
    fetch_public_policy,
    process_policy_job,
)
from prsa_control import ControlPlane


POLICY_TEXT = "Software must not set any deployment location to Russia or Russian territories."


def _proposal(excerpt=POLICY_TEXT):
    return {
        "control_id": "sanctions.russia-location",
        "title": "Prohibited deployment location",
        "description": "Detects prohibited Russian deployment locations.",
        "prohibited_condition": "Deployment locations must not target Russia.",
        "control_type": "literal_value",
        "severity": "ERROR",
        "scope": {"file_globs": ["*.yaml"]},
        "exclusions": ["Documentation", "Tests that verify rejection"],
        "clarification_questions": [],
        "source_reference": {"clause_id": "clause-00001", "paragraph": 1, "excerpt": excerpt},
        "confidence": 0.98,
        "match": {"prohibited_values": ["Russia", "ru-central", "Russian Federation"], "aliases": ["RU"], "file_globs": ["*.yaml"]},
        "tests": [
            {"name": "positive", "file": "deploy.yaml", "content": 'region: "ru-central"', "should_match": True},
            {"name": "negative", "file": "deploy.yaml", "content": 'region: "westeurope"', "should_match": False},
        ],
    }


def test_text_extraction_preserves_paragraph_reference():
    extraction = extract_document("policy.txt", (POLICY_TEXT + "\n\nApproved exceptions must be documented.").encode())
    assert extraction.clauses[0]["paragraph"] == 1
    assert extraction.clauses[0]["excerpt"] == POLICY_TEXT


def test_word_extraction_preserves_heading_and_table_rows():
    document = Document()
    document.add_heading("Geographic restrictions", level=1)
    document.add_paragraph(POLICY_TEXT)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Country"
    table.cell(0, 1).text = "RU"
    stream = BytesIO()
    document.save(stream)
    extraction = extract_document("policy.docx", stream.getvalue())
    assert extraction.clauses[0]["section"] == "Geographic restrictions"
    assert any(item["section"] == "Table 1" and "RU" in item["excerpt"] for item in extraction.clauses)


def test_image_only_pdf_requires_ocr_clarification():
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = BytesIO()
    writer.write(stream)
    with pytest.raises(PolicyEngineError, match="no extractable text"):
        extract_document("scanned.pdf", stream.getvalue())


def test_failed_ingestion_reports_an_actionable_reason_to_the_administrator():
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = BytesIO()
    writer.write(stream)
    plane = ControlPlane(connection_string="")
    policy = plane.save_policy_document(
        {"title": "Sanctions", "version": "2026-01", "filename": "scanned.pdf", "input_type": "upload"},
        stream.getvalue(),
        actor="author@example.com",
    )
    job = plane.record_policy_job(policy["document_id"], policy["version"], actor="author@example.com")

    with pytest.raises(PolicyEngineError):
        process_policy_job(job, plane, interpreter=FakeInterpreter())

    stored_job = plane.get_policy_job(job["job_id"])
    assert stored_job["status"] == "failed"
    assert stored_job["phase"] == "Policy ingestion failed"
    assert "provide an OCR-processed file" in stored_job["errors"][0]
    stored_policy = plane.get_policy(policy["document_id"], policy["version"])
    assert stored_policy["ingestion_status"] == "failed"
    assert stored_policy["status"] == "needs_clarification"


def test_compiler_rejects_invented_policy_citation():
    policy = {"document_id": "sanctions", "version": "1", "title": "Sanctions"}
    clauses = [{"clause_id": "clause-00001", "paragraph": 1, "excerpt": POLICY_TEXT}]
    with pytest.raises(PolicyEngineError, match="unverified"):
        compile_proposal(_proposal("This sentence was invented."), policy, clauses)


def test_compiler_runs_positive_and_negative_tests():
    policy = {"document_id": "sanctions", "version": "1", "title": "Sanctions"}
    clauses = [{"clause_id": "clause-00001", "paragraph": 1, "excerpt": POLICY_TEXT}]
    control = compile_proposal(_proposal(), policy, clauses)
    assert control["validation"]["passed"] is True
    assert [item["actual"] for item in control["validation"]["tests"]] == [True, False]


def test_compiler_normalizes_common_policy_risk_severity():
    policy = {"document_id": "sanctions", "version": "1", "title": "Sanctions"}
    clauses = [{"clause_id": "clause-00001", "paragraph": 1, "excerpt": POLICY_TEXT}]
    proposal = _proposal()
    proposal["severity"] = "high"

    control = compile_proposal(proposal, policy, clauses)

    assert control["severity"] == "ERROR"


def test_ast_compiler_forces_control_id_into_semgrep_rule_ids(monkeypatch):
    policy = {"document_id": "sanctions", "version": "1", "title": "Sanctions"}
    clauses = [{"clause_id": "clause-00001", "paragraph": 1, "excerpt": POLICY_TEXT}]
    proposal = _proposal()
    proposal["control_type"] = "ast"
    proposal["match"] = {"semgrep_yaml": "rules:\n  - id: arbitrary\n    languages: [python]\n    message: prohibited\n    severity: ERROR\n    pattern: bad(...)\n"}
    monkeypatch.setattr("policy_engine.run_semgrep", lambda files, **kwargs: [object()] if "ru-central" in next(iter(files.values())) else [])
    control = compile_proposal(proposal, policy, clauses)
    assert "id: sanctions.russia-location.arbitrary" in control["detector"]["semgrep_yaml"]


def test_url_ingestion_rejects_local_targets():
    with pytest.raises(PolicyEngineError):
        _validate_public_https("https://localhost/policy.pdf")
    with pytest.raises(PolicyEngineError):
        _validate_public_https("http://example.com/policy.pdf")


def test_url_fetch_enforces_media_type(monkeypatch):
    monkeypatch.setattr("policy_engine._validate_public_https", lambda url: None)
    client = httpx.Client(transport=httpx.MockTransport(lambda request: httpx.Response(200, headers={"content-type": "text/html"}, content=b"no")))
    with pytest.raises(PolicyEngineError, match="content type"):
        fetch_public_policy("https://example.com/policy", client=client)


def test_url_fetch_returns_bounded_supported_document(monkeypatch):
    monkeypatch.setattr("policy_engine._validate_public_https", lambda url: None)
    client = httpx.Client(transport=httpx.MockTransport(lambda request: httpx.Response(200, headers={"content-type": "text/plain"}, content=POLICY_TEXT.encode())))
    content, filename, media_type = fetch_public_policy("https://example.com/policy.txt", client=client)
    assert content.decode() == POLICY_TEXT
    assert filename == "policy.txt"
    assert media_type == "text/plain"


def test_policy_interpreter_batches_large_documents_without_duplicating_full_text():
    class Completions:
        def __init__(self):
            self.calls = []

        def create(self, **kwargs):
            self.calls.append(kwargs)
            message = type("Message", (), {"content": '{"controls":[],"obligations":[],"exceptions":[],"defined_terms":{}}'})()
            choice = type("Choice", (), {"message": message})()
            return type("Response", (), {"choices": [choice]})()

    completions = Completions()
    fake = type("Client", (), {"chat": type("Chat", (), {"completions": completions})()})()
    clauses = [{"clause_id": f"c-{index}", "paragraph": index, "excerpt": "x" * 6000} for index in range(40)]
    result = AzureOpenAIPolicyInterpreter(client=fake, deployment="model").interpret({"title": "Large"}, "ignored" * 100_000, clauses)
    assert result["controls"] == []
    assert len(completions.calls) == 2
    assert all('"text"' not in call["messages"][1]["content"] for call in completions.calls)
    assert all(call["response_format"]["type"] == "json_schema" for call in completions.calls)
    assert all(call["response_format"]["json_schema"]["strict"] is True for call in completions.calls)


def test_policy_interpreter_retries_empty_reasoning_completion_at_low_effort():
    class Completions:
        def __init__(self):
            self.calls = []

        def create(self, **kwargs):
            self.calls.append(kwargs)
            content = "" if len(self.calls) == 1 else '{"controls":[],"obligations":[],"exceptions":[],"defined_terms":{}}'
            message = type("Message", (), {"content": content})()
            choice = type("Choice", (), {"message": message, "finish_reason": "length" if not content else "stop"})()
            return type("Response", (), {"choices": [choice]})()

    completions = Completions()
    fake = type("Client", (), {"chat": type("Chat", (), {"completions": completions})()})()
    clauses = [{"clause_id": "c-1", "paragraph": 1, "excerpt": "A requirement."}]

    result = AzureOpenAIPolicyInterpreter(client=fake, deployment="model").interpret({"title": "Policy"}, "ignored", clauses)

    assert result["controls"] == []
    assert [call["reasoning_effort"] for call in completions.calls] == ["medium", "low"]


def test_policy_interpreter_keeps_obligation_ids_unique_across_batches():
    clauses = [
        {"clause_id": "clause-a", "paragraph": 1, "excerpt": "A" * 120_000},
        {"clause_id": "clause-b", "paragraph": 2, "excerpt": "B" * 120_000},
    ]

    class Completions:
        def __init__(self):
            self.index = 0

        def create(self, **kwargs):
            clause = clauses[self.index]
            self.index += 1
            payload = {
                "obligations": [{
                    "obligation_id": "requirement",
                    "statement": f"Requirement {self.index}",
                    "detection_surfaces": ["repository_settings"],
                    "source_reference": clause,
                }],
                "controls": [],
                "exceptions": [],
                "defined_terms": {},
            }
            message = type("Message", (), {"content": __import__("json").dumps(payload)})()
            return type("Response", (), {"choices": [type("Choice", (), {"message": message})()]})()

    completions = Completions()
    fake = type("Client", (), {"chat": type("Chat", (), {"completions": completions})()})()

    result = AzureOpenAIPolicyInterpreter(client=fake, deployment="model").interpret({"title": "Large"}, "ignored", clauses)

    assert [item["obligation_id"] for item in result["obligations"]] == ["requirement", "requirement-2"]
    assert {tuple(item["obligation_ids"]) for item in result["controls"]} == {("requirement",), ("requirement-2",)}


class FakeInterpreter:
    def interpret(self, policy, text, clauses):
        assert text == POLICY_TEXT
        return {"defined_terms": {}, "controls": [_proposal()]}


def test_policy_job_extracts_proposes_validates_and_persists():
    plane = ControlPlane(connection_string="")
    policy = plane.save_policy_document({"title": "Sanctions", "version": "2026-01", "filename": "policy.txt"}, POLICY_TEXT)
    job = plane.record_policy_job(policy["document_id"], policy["version"])
    controls = process_policy_job(job, plane, interpreter=FakeInterpreter())
    assert controls[0]["validation"]["passed"] is True
    assert plane.get_policy_job(job["job_id"])["status"] == "completed"
    assert plane.get_policy(policy["document_id"], policy["version"])["clause_count"] == 1
    assert plane.get_policy_analysis(policy["document_id"], policy["version"])["controls"][0]["control_id"] == "sanctions.russia-location"


def test_policy_job_retry_reuses_controls_and_clears_stale_errors():
    plane = ControlPlane(connection_string="")
    policy = plane.save_policy_document({"title": "Sanctions", "version": "2026-01", "filename": "policy.txt"}, POLICY_TEXT)
    job = plane.record_policy_job(policy["document_id"], policy["version"])

    first = process_policy_job(job, plane, interpreter=FakeInterpreter())
    plane.update_policy_job(job["job_id"], errors=["transient retry error"])
    retried = process_policy_job(job, plane, interpreter=FakeInterpreter())

    assert [(item["control_id"], item["version"]) for item in retried] == [(item["control_id"], item["version"]) for item in first]
    stored = plane.get_policy_job(job["job_id"])
    assert stored["status"] == "completed"
    assert stored["errors"] == []


def test_openai_model_version_policy_shape_errors_become_clarification_not_failure():
    text = "OpenAI usage: model prior to GPT 5.5 shall not be used"

    class MalformedInterpreter:
        def interpret(self, policy, extracted_text, clauses):
            assert extracted_text == text
            return {
                "obligations": [{
                    "obligation_id": "approved-openai-model",
                    "statement": "Do not use OpenAI models prior to GPT 5.5.",
                    "detection_surfaces": ["source_literals", "configuration_iac"],
                    "source_reference": text,
                }],
                "controls": [{
                    "control_id": "openai.minimum-model-version",
                    "obligation_ids": ["approved-openai-model"],
                    "title": "Minimum OpenAI model version",
                    "description": "Review selected OpenAI model versions.",
                    "prohibited_condition": "OpenAI model is older than GPT 5.5.",
                    "control_type": "manual_review",
                    "severity": "high",
                    "scope": "Evidence of model selection in source code and configuration",
                    "exclusions": [],
                    "clarification_questions": ["Provide the authoritative model catalog and ordering."],
                    "source_reference": text,
                    "confidence": "low",
                    "match": {},
                    "tests": [
                        {"name": "old model", "file": "policy-review.txt", "content": "gpt-4", "should_match": True},
                        {"name": "unrelated", "file": "policy-review.txt", "content": "hello", "should_match": False},
                    ],
                }],
            }

    plane = ControlPlane(connection_string="")
    policy = plane.save_policy_document(
        {"title": "OpenAI usage", "version": "1.0", "filename": "policy.txt"}, text
    )
    job = plane.record_policy_job(policy["document_id"], policy["version"])

    [control] = process_policy_job(job, plane, interpreter=MalformedInterpreter())

    assert control["confidence"] == 0.25
    assert control["severity"] == "ERROR"
    assert control["scope"] == {}
    assert control["state"] == "needs_clarification"
    assert any("scope was malformed" in item for item in control["clarification_questions"])
    assert control["source_reference"]["excerpt"] == text
    assert plane.get_policy_job(job["job_id"])["status"] == "completed"
    stored_policy = plane.get_policy(policy["document_id"], policy["version"])
    assert stored_policy["ingestion_status"] == "completed"
    assert stored_policy["status"] == "needs_clarification"
