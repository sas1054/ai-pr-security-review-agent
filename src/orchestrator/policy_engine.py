"""Natural-language policy ingestion with a deterministic control boundary."""

from __future__ import annotations

import ipaddress
import json
import os
import re
import socket
from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePosixPath
from typing import Any, Protocol
from urllib.parse import unquote, urljoin, urlparse

import httpx
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from openai import OpenAI

from prsa_control import ControlPlane
from scanner import Finding, run_semgrep, run_typed_control_scan


MAX_POLICY_BYTES = 20 * 1024 * 1024
ALLOWED_MEDIA_TYPES = {
    "text/plain": "txt",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


class PolicyEngineError(RuntimeError):
    pass


CONFIDENCE_LABELS = {
    "none": 0.0,
    "unknown": 0.0,
    "low": 0.25,
    "medium": 0.5,
    "moderate": 0.5,
    "high": 0.85,
    "very high": 0.95,
}
CONTROL_TYPE_PREFERENCE_LABELS = {
    "auto": "automatic selection based on the obligation and available evidence",
    "literal_value": "exact values or tokens such as model IDs, prohibited names, or environment values",
    "pattern": "bounded text patterns such as naming conventions or forbidden syntax",
    "ast": "code structure and call patterns validated through Semgrep",
    "config_iac": "structured configuration or infrastructure fields and values",
    "url_domain": "URLs, hostnames, and service domains",
    "dependency": "package and dependency manifests",
    "semantic_review": "behavior that requires model-assisted evidence from the changed code",
    "manual_review": "a human decision when no reliable automated detector is possible",
}

SOURCE_REFERENCE_SCHEMA = {
    "type": "object",
    "properties": {
        "clause_id": {"type": "string"},
        "page": {"type": ["integer", "null"]},
        "section": {"type": ["string", "null"]},
        "paragraph": {"type": ["integer", "null"]},
        "excerpt": {"type": "string"},
    },
    "required": ["clause_id", "page", "section", "paragraph", "excerpt"],
    "additionalProperties": False,
}

STRING_ARRAY_SCHEMA = {"type": "array", "items": {"type": "string"}}

POLICY_PROPOSAL_RESPONSE_FORMAT = {
    "type": "json_schema",
    "json_schema": {
        "name": "policy_control_proposal",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "controls": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "obligation_ids": STRING_ARRAY_SCHEMA,
                            "control_id": {"type": "string"},
                            "title": {"type": "string"},
                            "description": {"type": "string"},
                            "prohibited_condition": {"type": "string"},
                            "control_type": {"type": "string", "enum": ["literal_value", "pattern", "ast", "dependency", "url_domain", "config_iac", "semantic_review", "manual_review"]},
                            "severity": {"type": "string"},
                            "scope": {
                                "type": "object",
                                "properties": {key: STRING_ARRAY_SCHEMA for key in ("file_globs", "exclude_globs", "repositories", "branches", "environments")},
                                "required": ["file_globs", "exclude_globs", "repositories", "branches", "environments"],
                                "additionalProperties": False,
                            },
                            "exclusions": STRING_ARRAY_SCHEMA,
                            "clarification_questions": STRING_ARRAY_SCHEMA,
                            "source_reference": SOURCE_REFERENCE_SCHEMA,
                            "confidence": {"type": "number", "minimum": 0, "maximum": 1},
                            "match": {
                                "type": "object",
                                "properties": {
                                    **{key: STRING_ARRAY_SCHEMA for key in ("prohibited_values", "aliases", "field_names", "patterns", "packages", "package_prefixes", "domains", "file_globs", "exclude_globs")},
                                    "semgrep_yaml": {"type": "string"},
                                },
                                "required": ["prohibited_values", "aliases", "field_names", "patterns", "packages", "package_prefixes", "domains", "file_globs", "exclude_globs", "semgrep_yaml"],
                                "additionalProperties": False,
                            },
                            "detector_provenance": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {"value": {"type": "string"}, "source_kind": {"type": "string"}, "reference": {"type": "string"}},
                                    "required": ["value", "source_kind", "reference"],
                                    "additionalProperties": False,
                                },
                            },
                            "tests": {
                                "type": "array",
                                "items": {
                                    "type": "object",
                                    "properties": {"name": {"type": "string"}, "file": {"type": "string"}, "content": {"type": "string"}, "should_match": {"type": "boolean"}},
                                    "required": ["name", "file", "content", "should_match"],
                                    "additionalProperties": False,
                                },
                            },
                        },
                        "required": ["obligation_ids", "control_id", "title", "description", "prohibited_condition", "control_type", "severity", "scope", "exclusions", "clarification_questions", "source_reference", "confidence", "match", "detector_provenance", "tests"],
                        "additionalProperties": False,
                    },
                },
                "obligations": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "obligation_id": {"type": "string"},
                            "statement": {"type": "string"},
                            "source_reference": SOURCE_REFERENCE_SCHEMA,
                            "enforceability": {"type": "string"},
                            "detection_surfaces": STRING_ARRAY_SCHEMA,
                            "clarification_questions": STRING_ARRAY_SCHEMA,
                        },
                        "required": ["obligation_id", "statement", "source_reference", "enforceability", "detection_surfaces", "clarification_questions"],
                        "additionalProperties": False,
                    },
                },
                "exceptions": STRING_ARRAY_SCHEMA,
                "effective_dates": STRING_ARRAY_SCHEMA,
                "defined_terms": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {"term": {"type": "string"}, "definition": {"type": "string"}},
                        "required": ["term", "definition"],
                        "additionalProperties": False,
                    },
                },
                "document_scope": STRING_ARRAY_SCHEMA,
            },
            "required": ["controls", "obligations", "exceptions", "effective_dates", "defined_terms", "document_scope"],
            "additionalProperties": False,
        },
    },
}


def _list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _string_list(value: Any) -> list[str]:
    return [str(item).strip() for item in _list(value) if str(item).strip()]


def _normalize_confidence(value: Any) -> float:
    if isinstance(value, str):
        normalized = value.strip().lower().replace("_", " ").replace("-", " ")
        if normalized in CONFIDENCE_LABELS:
            return CONFIDENCE_LABELS[normalized]
    try:
        return min(1.0, max(0.0, float(value or 0)))
    except (TypeError, ValueError):
        return 0.0


def _recover_source(value: Any, clauses: list[dict[str, Any]]) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        excerpt = value.strip()
        for clause in clauses:
            if excerpt in str(clause.get("excerpt") or ""):
                return {**clause, "excerpt": excerpt}
    return dict(clauses[0]) if clauses else {}


def _normalize_model_proposal(proposal: Any, clauses: list[dict[str, Any]]) -> dict[str, Any]:
    """Make untrusted model JSON safe and turn shape defects into clarifications."""
    if not isinstance(proposal, dict):
        proposal = {}
    normalized = dict(proposal)
    obligations: list[dict[str, Any]] = []
    for item in _list(proposal.get("obligations")):
        raw = dict(item) if isinstance(item, dict) else {"statement": str(item)}
        questions = _string_list(raw.get("clarification_questions"))
        if not isinstance(raw.get("source_reference"), dict):
            questions.append("The generated obligation citation was malformed; confirm the cited policy clause.")
        raw["source_reference"] = _recover_source(raw.get("source_reference"), clauses)
        raw["detection_surfaces"] = _string_list(raw.get("detection_surfaces", raw.get("surfaces")))
        raw["clarification_questions"] = list(dict.fromkeys(questions))
        obligations.append(raw)
    controls: list[dict[str, Any]] = []
    for item in _list(proposal.get("controls")):
        if not isinstance(item, dict):
            continue
        raw = dict(item)
        questions = _string_list(raw.get("clarification_questions"))
        for field in ("scope", "match"):
            if not isinstance(raw.get(field), dict):
                if raw.get(field) not in (None, {}, ""):
                    questions.append(f"The generated {field} was malformed; define it before approving this control.")
                raw[field] = {}
        if not isinstance(raw.get("source_reference"), dict):
            questions.append("The generated control citation was malformed; confirm the cited policy clause.")
        raw["source_reference"] = _recover_source(raw.get("source_reference"), clauses)
        for field in ("obligation_ids", "exclusions", "detection_surfaces"):
            raw[field] = _string_list(raw.get(field))
        raw["detector_provenance"] = [value for value in _list(raw.get("detector_provenance")) if isinstance(value, dict)]
        raw["tests"] = [value for value in _list(raw.get("tests")) if isinstance(value, dict)]
        raw["confidence"] = _normalize_confidence(raw.get("confidence"))
        raw["clarification_questions"] = list(dict.fromkeys(questions))[:25]
        controls.append(raw)
    normalized["obligations"] = obligations
    normalized["controls"] = controls
    for field in ("exceptions", "effective_dates"):
        normalized[field] = _list(proposal.get(field))
    defined_terms = proposal.get("defined_terms")
    if isinstance(defined_terms, list):
        defined_terms = {
            str(item.get("term")): str(item.get("definition"))
            for item in defined_terms
            if isinstance(item, dict) and item.get("term")
        }
    normalized["defined_terms"] = defined_terms if isinstance(defined_terms, dict) else {}
    return normalized


@dataclass
class Extraction:
    text: str
    clauses: list[dict[str, Any]]


def _paragraphs(text: str) -> list[str]:
    return [value.strip() for value in re.split(r"\n\s*\n", text.replace("\r\n", "\n")) if value.strip()]


def _text_extraction(text: str, *, page: int | None = None, initial_section: str = "") -> Extraction:
    clauses: list[dict[str, Any]] = []
    section = initial_section
    pieces = _paragraphs(text)
    clause_index = 0
    for index, paragraph in enumerate(pieces, 1):
        first_line = paragraph.splitlines()[0].strip()
        if len(first_line) <= 100 and (first_line.endswith(":") or re.match(r"^\d+(?:\.\d+)*\s+", first_line)):
            section = first_line.rstrip(":")
        for offset in range(0, len(paragraph), 6000):
            clause_index += 1
            clause: dict[str, Any] = {
                "clause_id": f"clause-{clause_index:05d}",
                "paragraph": index,
                "paragraph_part": offset // 6000 + 1,
                "section": section,
                "excerpt": paragraph[offset : offset + 6000],
            }
            if page is not None:
                clause["page"] = page
            clauses.append(clause)
    return Extraction(text="\n\n".join(pieces), clauses=clauses)


def extract_document(filename: str, content: bytes) -> Extraction:
    extension = PurePosixPath(filename.lower()).suffix
    if len(content) > MAX_POLICY_BYTES:
        raise PolicyEngineError("Policy document exceeds the 20 MB limit")
    if extension == ".txt":
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError:
            text = content.decode("utf-8", errors="replace")
        return _text_extraction(text)
    if extension == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - deployment dependency
            raise PolicyEngineError("PDF extraction dependency is unavailable") from exc
        reader = PdfReader(BytesIO(content))
        all_text: list[str] = []
        clauses: list[dict[str, Any]] = []
        for page_number, page in enumerate(reader.pages, 1):
            page_text = page.extract_text() or ""
            extraction = _text_extraction(page_text, page=page_number)
            all_text.append(extraction.text)
            for clause_number, clause in enumerate(extraction.clauses, 1):
                clause["clause_id"] = f"page-{page_number}-{clause_number:05d}"
                clauses.append(clause)
        text = "\n\n".join(item for item in all_text if item)
        if not text.strip():
            raise PolicyEngineError("PDF contains no extractable text; provide an OCR-processed file")
        return Extraction(text=text, clauses=clauses)
    if extension == ".docx":
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - deployment dependency
            raise PolicyEngineError("Word extraction dependency is unavailable") from exc
        document = Document(BytesIO(content))
        section = ""
        clauses: list[dict[str, Any]] = []
        text_parts: list[str] = []
        paragraph_number = 0
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if not value:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                section = value
                continue
            paragraph_number += 1
            text_parts.append(value)
            for offset in range(0, len(value), 6000):
                clauses.append(
                    {
                        "clause_id": f"paragraph-{paragraph_number:05d}-{offset // 6000 + 1}",
                        "paragraph": paragraph_number,
                        "paragraph_part": offset // 6000 + 1,
                        "section": section,
                        "excerpt": value[offset : offset + 6000],
                    }
                )
        for table_number, table in enumerate(document.tables, 1):
            for row_number, row in enumerate(table.rows, 1):
                value = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if not value:
                    continue
                paragraph_number += 1
                text_parts.append(value)
                for offset in range(0, len(value), 6000):
                    clauses.append(
                        {
                            "clause_id": f"table-{table_number:03d}-row-{row_number:05d}-{offset // 6000 + 1}",
                            "paragraph": paragraph_number,
                            "paragraph_part": offset // 6000 + 1,
                            "section": f"Table {table_number}",
                            "excerpt": value[offset : offset + 6000],
                        }
                    )
        if not text_parts:
            raise PolicyEngineError("Word document contains no extractable text")
        return Extraction(text="\n\n".join(text_parts), clauses=clauses)
    raise PolicyEngineError("Only PDF, DOCX, and TXT policy documents are supported")


def _validate_public_https(url: str) -> None:
    parsed = urlparse(url)
    if parsed.scheme.lower() != "https" or not parsed.hostname or parsed.username or parsed.password:
        raise PolicyEngineError("Policy URLs must be unauthenticated HTTPS URLs")
    if parsed.port not in (None, 443):
        raise PolicyEngineError("Policy URLs must use the standard HTTPS port")
    try:
        addresses = socket.getaddrinfo(parsed.hostname, 443, type=socket.SOCK_STREAM)
    except socket.gaierror as exc:
        raise PolicyEngineError("Policy URL host could not be resolved") from exc
    if not addresses:
        raise PolicyEngineError("Policy URL host has no addresses")
    for address in addresses:
        ip = ipaddress.ip_address(address[4][0])
        if not ip.is_global:
            raise PolicyEngineError("Policy URL resolves to a private, local, or reserved address")


def fetch_public_policy(url: str, *, client: httpx.Client | None = None) -> tuple[bytes, str, str]:
    owned_client = client is None
    session = client or httpx.Client(timeout=20, follow_redirects=False)
    current = url
    try:
        for _ in range(4):
            _validate_public_https(current)
            with session.stream("GET", current, headers={"Accept": ", ".join(ALLOWED_MEDIA_TYPES)}) as response:
                if response.status_code in {301, 302, 303, 307, 308}:
                    location = response.headers.get("location")
                    if not location:
                        raise PolicyEngineError("Policy URL redirect has no location")
                    current = urljoin(current, location)
                    continue
                response.raise_for_status()
                media_type = response.headers.get("content-type", "").split(";", 1)[0].strip().lower()
                if media_type not in ALLOWED_MEDIA_TYPES:
                    raise PolicyEngineError(f"Unsupported policy URL content type: {media_type or 'missing'}")
                content_length = int(response.headers.get("content-length") or 0)
                if content_length > MAX_POLICY_BYTES:
                    raise PolicyEngineError("Policy URL exceeds the 20 MB limit")
                chunks: list[bytes] = []
                total = 0
                for chunk in response.iter_bytes():
                    total += len(chunk)
                    if total > MAX_POLICY_BYTES:
                        raise PolicyEngineError("Policy URL exceeds the 20 MB limit")
                    chunks.append(chunk)
                basename = unquote(PurePosixPath(urlparse(current).path).name)
                expected_extension = ALLOWED_MEDIA_TYPES[media_type]
                filename = basename if basename.lower().endswith("." + expected_extension) else f"source.{expected_extension}"
                return b"".join(chunks), filename, media_type
        raise PolicyEngineError("Policy URL exceeded the three-redirect limit")
    except httpx.HTTPError as exc:
        raise PolicyEngineError(f"Could not fetch policy URL: {exc}") from exc
    finally:
        if owned_client:
            session.close()


class PolicyInterpreter(Protocol):
    def interpret(self, policy: dict[str, Any], text: str, clauses: list[dict[str, Any]]) -> dict[str, Any]: ...


CANONICAL_SURFACES = {
    "source_literals",
    "code_structure",
    "dependencies",
    "service_endpoints",
    "configuration_iac",
    "semantic_behavior",
    "repository_settings",
    "manual_evidence",
}

CONTROL_TYPE_SURFACES = {
    "literal_value": {"source_literals"},
    "pattern": {"code_structure"},
    "ast": {"code_structure"},
    "dependency": {"dependencies"},
    "url_domain": {"service_endpoints"},
    "config_iac": {"configuration_iac"},
    "semantic_review": {"semantic_behavior"},
    "manual_review": {"manual_evidence"},
}

DETECTOR_TERM_KEYS = {"prohibited_values", "aliases", "field_names", "packages", "package_prefixes", "domains"}


def _canonical_surface(value: Any) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "_", str(value or "").strip().lower()).strip("_")
    aliases = {
        "literal": "source_literals", "literals": "source_literals", "string_literals": "source_literals",
        "source_code": "code_structure", "ast": "code_structure", "imports": "code_structure", "function_calls": "code_structure",
        "dependency": "dependencies", "dependency_manifest": "dependencies", "dependency_manifests": "dependencies", "lock_files": "dependencies",
        "url": "service_endpoints", "urls": "service_endpoints", "domain": "service_endpoints", "domains": "service_endpoints",
        "api_calls": "service_endpoints", "api_endpoints": "service_endpoints", "network_endpoints": "service_endpoints",
        "configuration": "configuration_iac", "config": "configuration_iac", "iac": "configuration_iac",
        "deployment": "configuration_iac", "deployment_metadata": "configuration_iac", "infrastructure": "configuration_iac",
        "semantic": "semantic_behavior", "behavior": "semantic_behavior", "data_flow": "semantic_behavior",
        "repo_settings": "repository_settings", "repository_configuration": "repository_settings",
        "manual": "manual_evidence", "human_review": "manual_evidence",
    }
    return aliases.get(normalized, normalized)


def _source_for_obligation(raw: dict[str, Any], clauses: list[dict[str, Any]]) -> dict[str, Any] | None:
    try:
        return _verified_source(raw, clauses)
    except PolicyEngineError:
        return None


def assess_policy_proposal(proposal: dict[str, Any], clauses: list[dict[str, Any]]) -> dict[str, Any]:
    """Build a generic obligation-to-control coverage matrix and fail ambiguity closed.

    This does not decide policy meaning. It checks that the model accounted for every
    obligation and declared scan surface, and that concrete detector vocabulary is
    either quoted by the policy or explicitly sent back for human clarification.
    """
    proposal = _normalize_model_proposal(proposal, clauses)
    controls = [dict(item) for item in proposal.get("controls", []) if isinstance(item, dict)]
    obligations: list[dict[str, Any]] = []
    for index, item in enumerate(proposal.get("obligations", []), 1):
        raw = dict(item) if isinstance(item, dict) else {"statement": str(item)}
        obligation_id = re.sub(r"[^a-z0-9._-]+", "-", str(raw.get("obligation_id") or f"obligation-{index:03d}").lower()).strip("-")
        surfaces = sorted(
            {
                _canonical_surface(value)
                for value in _list(raw.get("detection_surfaces", raw.get("surfaces", [])))
                if _canonical_surface(value) in CANONICAL_SURFACES
            }
        )
        source = _source_for_obligation(raw, clauses)
        obligations.append(
            {
                **raw,
                "obligation_id": obligation_id,
                "statement": str(raw.get("statement") or raw.get("obligation") or "").strip(),
                "detection_surfaces": surfaces,
                "source_reference": source or raw.get("source_reference") or {},
            }
        )
    if not obligations:
        # Older model responses used string obligations or omitted them. Derive an
        # explicit obligation per cited control so coverage is visible, not implicit.
        for index, control in enumerate(controls, 1):
            obligations.append(
                {
                    "obligation_id": f"obligation-{index:03d}",
                    "statement": str(control.get("prohibited_condition") or control.get("description") or "").strip(),
                    "detection_surfaces": sorted(CONTROL_TYPE_SURFACES.get(str(control.get("control_type") or ""), set())),
                    "source_reference": control.get("source_reference") or {},
                    "derived_from_control": True,
                }
            )

    obligation_ids = {item["obligation_id"] for item in obligations}
    by_obligation: dict[str, list[dict[str, Any]]] = {item: [] for item in obligation_ids}
    for index, control in enumerate(controls):
        linked = [str(item) for item in _list(control.get("obligation_ids")) if str(item) in obligation_ids]
        if not linked and len(obligations) == 1:
            linked = [obligations[0]["obligation_id"]]
        if not linked:
            clause_id = str((control.get("source_reference") or {}).get("clause_id") or "")
            linked = [
                item["obligation_id"]
                for item in obligations
                if clause_id and clause_id == str((item.get("source_reference") or {}).get("clause_id") or "")
            ]
        control["obligation_ids"] = sorted(set(linked))
        control["detection_surfaces"] = sorted(CONTROL_TYPE_SURFACES.get(str(control.get("control_type") or ""), set()))
        questions = _string_list(control.get("clarification_questions"))
        source_reference = control.get("source_reference") if isinstance(control.get("source_reference"), dict) else {}
        source_text = str(source_reference.get("excerpt") or "").casefold()
        source_clause_id = str(source_reference.get("clause_id") or "")
        source_text += " " + " ".join(
            str(item.get("excerpt") or "").casefold()
            for item in clauses
            if source_clause_id and str(item.get("clause_id") or "") == source_clause_id
        )
        match = control.get("match") if isinstance(control.get("match"), dict) else {}
        for key in DETECTOR_TERM_KEYS:
            values = match.get(key, [])
            if not isinstance(values, list):
                continue
            for value in values:
                term = str(value).strip()
                if not term or term.casefold() in source_text:
                    continue
                questions.append(
                    f"Detector term '{term}' is not stated in the cited policy text. Provide an approved source or remove it."
                )
        control["clarification_questions"] = list(dict.fromkeys(questions))[:25]
        controls[index] = control
        for obligation_id in control["obligation_ids"]:
            by_obligation[obligation_id].append(control)

    # Never silently drop an extracted obligation. A cited, non-executable
    # obligation becomes a visible clarification placeholder that can never be
    # approved or activated in place.
    for obligation in obligations:
        obligation_id = obligation["obligation_id"]
        if by_obligation[obligation_id]:
            if not obligation.get("source_reference"):
                obligation["source_reference"] = by_obligation[obligation_id][0].get("source_reference") or {}
            continue
        source = _source_for_obligation(obligation, clauses)
        if not source:
            continue
        placeholder = {
            "control_id": f"{obligation_id}.coverage-review",
            "title": f"Clarify implementation for {obligation_id}",
            "description": "No reliable machine-executable PR control was generated for this obligation.",
            "prohibited_condition": obligation.get("statement") or "Policy obligation requires implementation planning.",
            "control_type": "manual_review",
            "severity": "WARNING",
            "scope": {},
            "exclusions": [],
            "clarification_questions": [
                f"No executable control implements obligation '{obligation_id}'. Define its PR scope and approved detection strategy."
            ],
            "source_reference": source,
            "confidence": 0.0,
            "match": {},
            "tests": [
                {"name": "requires implementation", "file": "policy-review.txt", "content": obligation.get("statement") or obligation_id, "should_match": True},
                {"name": "unrelated change", "file": "policy-review.txt", "content": "unrelated change", "should_match": False},
            ],
            "obligation_ids": [obligation_id],
            "detection_surfaces": ["manual_evidence"],
            "generated_coverage_placeholder": True,
        }
        controls.append(placeholder)
        by_obligation[obligation_id].append(placeholder)

    coverage: list[dict[str, Any]] = []
    policy_questions: list[str] = []
    for obligation in obligations:
        obligation_id = obligation["obligation_id"]
        linked = by_obligation[obligation_id]
        covered = sorted({surface for control in linked for surface in control.get("detection_surfaces", [])})
        expected = obligation.get("detection_surfaces") or []
        uncovered = sorted(set(expected) - set(covered))
        obligation_questions: list[str] = []
        if not expected:
            obligation_questions.append(f"Define the PR detection surfaces for obligation '{obligation_id}'.")
        if not linked:
            obligation_questions.append(f"No proposed control implements obligation '{obligation_id}'.")
        if uncovered:
            obligation_questions.append(
                f"Obligation '{obligation_id}' has no proposed control for: {', '.join(uncovered)}."
            )
        policy_questions.extend(obligation_questions)
        for question in obligation_questions:
            for target in linked[:1]:
                target["clarification_questions"] = list(dict.fromkeys([*target.get("clarification_questions", []), question]))[:25]
        coverage.append(
            {
                "obligation_id": obligation_id,
                "expected_surfaces": expected,
                "covered_surfaces": covered,
                "uncovered_surfaces": uncovered,
                "control_ids": [str(item.get("control_id") or "") for item in linked],
            }
        )
    proposal = {
        **proposal,
        "obligations": obligations,
        "controls": controls,
        "coverage": coverage,
        "coverage_questions": list(dict.fromkeys(policy_questions)),
    }
    proposal["coverage_complete"] = not proposal["coverage_questions"] and all(
        not item.get("clarification_questions") for item in controls
    )
    return proposal


class AzureOpenAIPolicyInterpreter:
    def __init__(self, *, client: Any | None = None, deployment: str | None = None):
        self.deployment = deployment or os.environ.get("AZURE_OPENAI_DEPLOYMENT", "")
        if client is not None:
            self.client = client
            return
        endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT", "")
        if not endpoint or not self.deployment:
            raise PolicyEngineError("AZURE_OPENAI_ENDPOINT and AZURE_OPENAI_DEPLOYMENT are required")
        api_key = os.environ.get("AZURE_OPENAI_API_KEY", "")
        if api_key and os.environ.get("HACKATHON_MODE", "").lower() != "true":
            raise PolicyEngineError("Azure OpenAI API keys are permitted only in hackathon mode")
        credential = api_key or get_bearer_token_provider(DefaultAzureCredential(), "https://ai.azure.com/.default")
        self.client = OpenAI(base_url=f"{endpoint.rstrip('/')}/openai/v1/", api_key=credential, timeout=90)

    def interpret(self, policy: dict[str, Any], text: str, clauses: list[dict[str, Any]]) -> dict[str, Any]:
        clause_payload = [
            {key: item.get(key) for key in ("clause_id", "page", "section", "paragraph", "excerpt") if item.get(key) not in (None, "")}
            for item in clauses
        ]
        preference = str(policy.get("control_type_preference") or "auto").strip().lower()
        if preference not in CONTROL_TYPE_PREFERENCE_LABELS:
            preference = "auto"
        preference_instruction = (
            f"The administrator selected '{preference}' ({CONTROL_TYPE_PREFERENCE_LABELS[preference]}). "
            "Prefer this control type whenever it is compatible with the obligation. "
            "Do not invent detector vocabulary or force an incompatible type; ask a clarification only for missing policy evidence, scope, or an unsupported detection surface. "
            if preference != "auto"
            else "The administrator selected automatic control-type selection; prefer the most specific deterministic type that the policy evidence supports. "
        )
        prompt = (
            "Build a policy-agnostic PR control plan from the supplied clauses. Return JSON with controls, obligations, exceptions, "
            "effective_dates, defined_terms, and document_scope. Each obligation must be an object with obligation_id, statement, "
            "source_reference, enforceability, and detection_surfaces. detection_surfaces may only use source_literals, code_structure, "
            "dependencies, service_endpoints, configuration_iac, semantic_behavior, repository_settings, or manual_evidence. "
            "Account for every enforceable obligation; do not optimize for any example domain. Each control needs obligation_ids, "
            "control_id, title, description, "
            "prohibited_condition, control_type, severity, scope, exclusions, clarification_questions, source_reference, confidence, "
            "match, detector_provenance, and tests. control_type must be literal_value, pattern, ast, dependency, url_domain, config_iac, semantic_review, "
            "or manual_review. source_reference must copy one supplied clause excerpt exactly and retain its clause/page/section/paragraph. "
            "match may contain prohibited_values, aliases, field_names, patterns, packages, package_prefixes, domains, file_globs, "
            "exclude_globs, or semgrep_yaml. A Semgrep rule ID must be stable; the compiler will prefix it with the control ID. tests "
            "contain file, content, and should_match, with at least one positive and one negative test per control. Never invent detector "
            "values, package names, domains, aliases, or field names. If a necessary term is absent from the cited clause, do not guess: "
            "add a clarification question requesting an approved catalog or source. detector_provenance lists only terms directly supported "
            "by policy text, with value, source_kind='policy', and reference=clause_id. Ambiguous scope or an uncovered detection surface "
            "must produce clarification_questions. Never invent a citation or silently broaden an obligation. Prefer deterministic control "
            "types; use semantic_review or manual_review when reliable compilation is impossible, and never imply those establish compliance. "
            + preference_instruction
        )
        # Output, not input context, is the limiting factor for policy plans:
        # each clause can result in obligations, controls, citations, and tests.
        # Keep batches deliberately small so a moderately sized policy cannot
        # consume the whole completion budget in a single response.
        max_batch_chars = 60_000
        max_batch_clauses = 12
        batches: list[list[dict[str, Any]]] = []
        current: list[dict[str, Any]] = []
        current_size = 0
        for clause in clause_payload:
            size = len(json.dumps(clause, ensure_ascii=False))
            if current and (len(current) >= max_batch_clauses or current_size + size > max_batch_chars):
                batches.append(current)
                current, current_size = [], 0
            current.append(clause)
            current_size += size
        if current:
            batches.append(current)
        if not batches:
            raise PolicyEngineError("Policy contains no extractable clauses")

        results: list[dict[str, Any]] = []
        for batch_number, batch in enumerate(batches, 1):
            content = ""
            finish_reason = "unknown"
            configured_effort = os.environ.get("LLM_REASONING_EFFORT", "medium")
            for effort in dict.fromkeys((configured_effort, "low")):
                response = self.client.chat.completions.create(
                    model=self.deployment,
                    max_completion_tokens=min(int(os.environ.get("LLM_MAX_OUTPUT_TOKENS", "16000")), 16000),
                    reasoning_effort=effort,
                    response_format=POLICY_PROPOSAL_RESPONSE_FORMAT,
                    messages=[
                        {"role": "developer", "content": prompt},
                        {
                            "role": "user",
                            "content": json.dumps(
                                {"policy": policy, "batch": batch_number, "batch_count": len(batches), "clauses": batch},
                                ensure_ascii=False,
                            ),
                        },
                    ],
                )
                content = response.choices[0].message.content or ""
                finish_reason = str(getattr(response.choices[0], "finish_reason", "unknown"))
                if content:
                    break
            if not content:
                raise PolicyEngineError(
                    f"Azure OpenAI returned an empty policy proposal after bounded retry (finish_reason={finish_reason})"
                )
            try:
                result = json.loads(content)
            except json.JSONDecodeError as exc:
                raise PolicyEngineError("Azure OpenAI returned invalid policy JSON") from exc
            if not isinstance(result, dict) or not isinstance(result.get("controls"), list):
                raise PolicyEngineError("Policy proposal must contain a controls array")
            results.append(_normalize_model_proposal(result, clauses))

        merged: dict[str, Any] = {
            "controls": [], "obligations": [], "exceptions": [], "effective_dates": [], "defined_terms": {}, "document_scope": []
        }
        seen_controls: dict[str, str] = {}
        seen_obligations: dict[str, str] = {}
        for result in results:
            obligation_id_map: dict[str, str] = {}
            if isinstance(result.get("obligations"), list):
                for raw_obligation in result["obligations"]:
                    if not isinstance(raw_obligation, dict):
                        merged["obligations"].append(raw_obligation)
                        continue
                    obligation = dict(raw_obligation)
                    original_id = str(obligation.get("obligation_id") or "")
                    statement = str(obligation.get("statement") or obligation.get("obligation") or "")
                    resolved_id = original_id
                    if original_id and original_id in seen_obligations and seen_obligations[original_id] != statement:
                        suffix = 2
                        while f"{original_id}-{suffix}" in seen_obligations:
                            suffix += 1
                        resolved_id = f"{original_id}-{suffix}"
                        obligation["obligation_id"] = resolved_id
                    if resolved_id:
                        obligation_id_map[original_id] = resolved_id
                        seen_obligations[resolved_id] = statement
                    if original_id and original_id in seen_obligations and seen_obligations[original_id] == statement and resolved_id == original_id:
                        if any(
                            isinstance(item, dict) and item.get("obligation_id") == original_id
                            for item in merged["obligations"]
                        ):
                            continue
                    merged["obligations"].append(obligation)
            for key in ("exceptions", "effective_dates"):
                if isinstance(result.get(key), list):
                    merged[key].extend(result[key])
            if isinstance(result.get("defined_terms"), dict):
                merged["defined_terms"].update(result["defined_terms"])
            if result.get("document_scope"):
                merged["document_scope"].append(result["document_scope"])
            for control in result["controls"]:
                if not isinstance(control, dict):
                    continue
                control = dict(control)
                if isinstance(control.get("obligation_ids"), list):
                    control["obligation_ids"] = [
                        obligation_id_map.get(str(item), str(item)) for item in control["obligation_ids"]
                    ]
                base = str(control.get("control_id") or control.get("title") or "control")
                condition = str(control.get("prohibited_condition") or "")
                if base in seen_controls and seen_controls[base] == condition:
                    existing = next(item for item in merged["controls"] if str(item.get("control_id") or item.get("title") or "control") == base)
                    existing_clause = str((existing.get("source_reference") or {}).get("clause_id") or "")
                    control_clause = str((control.get("source_reference") or {}).get("clause_id") or "")
                    if existing_clause and existing_clause == control_clause:
                        existing["obligation_ids"] = sorted(
                            set(str(item) for item in existing.get("obligation_ids", []))
                            | set(str(item) for item in control.get("obligation_ids", []))
                        )
                        continue
                if base in seen_controls:
                    suffix = 2
                    while f"{base}-{suffix}" in seen_controls:
                        suffix += 1
                    control = {**control, "control_id": f"{base}-{suffix}"}
                    base = str(control["control_id"])
                seen_controls[base] = condition
                merged["controls"].append(control)
        return assess_policy_proposal(merged, clauses)


class AzureOpenAISemanticControlScanner:
    """Produces review evidence only; it never returns or implies a compliance verdict."""

    def __init__(self, *, client: Any | None = None, deployment: str | None = None):
        interpreter = AzureOpenAIPolicyInterpreter(client=client, deployment=deployment)
        self.client = interpreter.client
        self.deployment = interpreter.deployment

    def scan(self, files: dict[str, str], controls: list[dict[str, Any]]) -> list[Finding]:
        candidates = [item for item in controls if item.get("control_type") in {"semantic_review", "manual_review"}]
        if not candidates or not files:
            return []
        bounded_files: dict[str, str] = {}
        remaining_chars = 300_000
        for path, content in files.items():
            if remaining_chars <= 0:
                break
            value = content[: min(100_000, remaining_chars)]
            bounded_files[path] = value
            remaining_chars -= len(value)
        request = {
            "controls": [
                {
                    "control_id": item.get("control_id"),
                    "version": item.get("version"),
                    "title": item.get("title"),
                    "prohibited_condition": item.get("prohibited_condition"),
                    "scope": item.get("scope"),
                    "exclusions": item.get("exclusions"),
                }
                for item in candidates
            ],
            "changed_files": bounded_files,
        }
        response = self.client.chat.completions.create(
            model=self.deployment,
            max_completion_tokens=min(int(os.environ.get("LLM_MAX_OUTPUT_TOKENS", "8000")), 8000),
            reasoning_effort=os.environ.get("LLM_REASONING_EFFORT", "medium"),
            response_format={"type": "json_object"},
            messages=[
                {
                    "role": "developer",
                    "content": (
                        "Identify code evidence that requires human review under the supplied controls. Return JSON with findings only. "
                        "Each finding needs control_id, file, line, matched_value, reason, and confidence. Use only supplied control IDs "
                        "and files. Never claim compliance, never suppress a deterministic finding, and omit uncertain speculation below 0.5 confidence."
                    ),
                },
                {"role": "user", "content": json.dumps(request, ensure_ascii=False)},
            ],
        )
        content = response.choices[0].message.content
        try:
            payload = json.loads(content or "{}")
        except json.JSONDecodeError as exc:
            raise PolicyEngineError("Semantic control scan returned invalid JSON") from exc
        by_id = {str(item.get("control_id")): item for item in candidates}
        findings: list[Finding] = []
        for raw in payload.get("findings", []):
            if not isinstance(raw, dict):
                continue
            control = by_id.get(str(raw.get("control_id") or ""))
            file_path = str(raw.get("file") or "")
            if not control or file_path not in bounded_files:
                continue
            try:
                confidence = float(raw.get("confidence") or 0)
                line = int(raw.get("line") or 0)
            except (TypeError, ValueError):
                continue
            if confidence < 0.5 or line < 1 or line > max(1, len(bounded_files[file_path].splitlines())):
                continue
            finding = Finding(
                tool="policy-semantic-review",
                rule_id=str(control["control_id"]),
                file=file_path if file_path.startswith("/") else f"/{file_path}",
                line=line,
                severity="WARNING",
                message=f"Human review required: {control.get('title', 'policy control')}",
                fix_hint=str(control.get("fix_hint") or "Review the evidence and document an approved exception if applicable."),
                control_id=str(control["control_id"]),
                control_version=str(control.get("version") or ""),
                reason=str(raw.get("reason") or control.get("prohibited_condition") or ""),
                policy_document=str(control.get("policy_title") or control.get("policy_document_id") or ""),
                policy_version=str(control.get("policy_version") or ""),
                source_reference=dict(control.get("source_reference") or {}),
                confidence=confidence,
                matched_value=str(raw.get("matched_value") or "semantic evidence"),
                review_required=True,
            )
            findings.append(finding)
        return findings


def _verified_source(raw: dict[str, Any], clauses: list[dict[str, Any]]) -> dict[str, Any]:
    source = raw.get("source_reference")
    if not isinstance(source, dict):
        raise PolicyEngineError("Generated control contains a malformed policy citation")
    excerpt = str(source.get("excerpt") or "").strip()
    clause_id = str(source.get("clause_id") or "")
    for clause in clauses:
        if clause_id and clause_id != str(clause.get("clause_id") or ""):
            continue
        actual = str(clause.get("excerpt") or "")
        if excerpt and excerpt in actual:
            return {**clause, "excerpt": excerpt}
    raise PolicyEngineError("Generated control contains an unverified policy citation")


def compile_proposal(raw: dict[str, Any], policy: dict[str, Any], clauses: list[dict[str, Any]]) -> dict[str, Any]:
    kind = str(raw.get("control_type") or "manual_review")
    control_id = re.sub(r"[^a-z0-9._-]+", "-", str(raw.get("control_id") or raw.get("title") or "control").strip().lower()).strip("._-")[:120] or "control"
    # Model-generated identifiers such as "ctl-001" are only meaningful
    # within a policy.  Namespace them so two policies cannot collide.
    policy_prefix = f"{policy['document_id']}."
    if not control_id.startswith(policy_prefix):
        control_id = f"{policy_prefix}{control_id}"[:120]
    match = raw.get("match") if isinstance(raw.get("match"), dict) else {}
    detector: dict[str, Any] = {
        key: value
        for key, value in match.items()
        if key in {"prohibited_values", "aliases", "field_names", "patterns", "packages", "package_prefixes", "domains", "file_globs", "exclude_globs", "semgrep_yaml"}
    }
    tests = [item for item in _list(raw.get("tests")) if isinstance(item, dict)]
    questions = _string_list(raw.get("clarification_questions"))
    source = _verified_source(raw, clauses)
    if kind == "ast" and detector.get("semgrep_yaml"):
        try:
            import yaml

            semgrep_config = yaml.safe_load(str(detector["semgrep_yaml"]))
            rules = semgrep_config.get("rules") if isinstance(semgrep_config, dict) else None
            if not isinstance(rules, list) or not rules:
                raise ValueError("rules are required")
            for index, rule in enumerate(rules):
                if not isinstance(rule, dict):
                    raise ValueError("rule must be an object")
                original = re.sub(r"[^a-zA-Z0-9_.-]+", "-", str(rule.get("id") or index))
                rule["id"] = f"{control_id}.{original}"
            detector["semgrep_yaml"] = yaml.safe_dump(semgrep_config, sort_keys=False)
        except Exception as exc:
            raise PolicyEngineError(f"AST control contains invalid Semgrep YAML: {exc}") from exc
    control = {
        "control_id": control_id,
        "version": str(raw.get("version") or "1.0"),
        "title": str(raw.get("title") or "Policy control"),
        "description": str(raw.get("description") or ""),
        "prohibited_condition": str(raw.get("prohibited_condition") or ""),
        "control_type": kind,
        "obligation_ids": _string_list(raw.get("obligation_ids")),
        "detection_surfaces": _string_list(raw.get("detection_surfaces")),
        "detector_provenance": [item for item in _list(raw.get("detector_provenance")) if isinstance(item, dict)],
        "generated_coverage_placeholder": bool(raw.get("generated_coverage_placeholder")),
        "severity": _normalize_severity(raw.get("severity")),
        "scope": raw.get("scope") if isinstance(raw.get("scope"), dict) else {},
        "examples": {
            "positive": [item.get("content", "") for item in tests if item.get("should_match") is True],
            "negative": [item.get("content", "") for item in tests if item.get("should_match") is False],
        },
        "exclusions": _string_list(raw.get("exclusions")),
        "clarification_questions": questions,
        "policy_document_id": policy["document_id"],
        "policy_version": policy["version"],
        "policy_title": policy["title"],
        "source_reference": source,
        "detector": detector,
        "confidence": _normalize_confidence(raw.get("confidence")),
        "fix_hint": str(raw.get("fix_hint") or "Use an approved alternative or request a documented exception."),
        "state": "needs_clarification" if questions else "draft",
    }
    control["validation"] = validate_compiled_control(control, tests)
    return control


def _normalize_severity(value: Any) -> str:
    """Translate common policy-risk labels into the scanner's severities."""
    normalized = str(value or "WARNING").strip().upper()
    if normalized in {"ERROR", "WARNING", "INFO"}:
        return normalized
    if normalized in {"BLOCKER", "CRITICAL", "HIGH", "SEV0", "SEV1"}:
        return "ERROR"
    if normalized in {"MEDIUM", "MODERATE", "WARN", "SEV2"}:
        return "WARNING"
    if normalized in {"LOW", "INFORMATIONAL", "NOTE", "SEV3", "SEV4"}:
        return "INFO"
    return "WARNING"


def validate_compiled_control(control: dict[str, Any], tests: list[dict[str, Any]]) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    if not tests:
        return {"passed": False, "tests": [], "error": "At least one positive and one negative test are required"}
    expectations = {bool(item.get("should_match")) for item in tests}
    if expectations != {False, True}:
        return {"passed": False, "tests": [], "error": "Positive and negative tests are both required"}
    for index, test in enumerate(tests):
        path = str(test.get("file") or "sample.txt")
        content = str(test.get("content") or "")
        expected = bool(test.get("should_match"))
        try:
            if control.get("control_type") == "ast":
                semgrep_yaml = str((control.get("detector") or {}).get("semgrep_yaml") or "")
                if not semgrep_yaml or len(semgrep_yaml) > 100_000:
                    raise PolicyEngineError("AST control requires bounded Semgrep YAML")
                actual = bool(run_semgrep({path: content}, policy_rules=[{"rule_id": control["control_id"], "semgrep_yaml": semgrep_yaml}]))
            elif control.get("control_type") in {"manual_review", "semantic_review"}:
                actual = expected
            else:
                actual = bool(run_typed_control_scan({path: content}, [control]))
            passed = actual == expected
            results.append({"name": str(test.get("name") or f"test-{index + 1}"), "expected": expected, "actual": actual, "passed": passed})
        except Exception as exc:
            results.append({"name": str(test.get("name") or f"test-{index + 1}"), "expected": expected, "actual": None, "passed": False, "error": str(exc)[:500]})
    return {"passed": all(item["passed"] for item in results), "tests": results}


def process_policy_job(job: dict[str, Any], controls: ControlPlane, interpreter: PolicyInterpreter | None = None) -> list[dict[str, Any]]:
    job_id = str(job.get("job_id") or "")
    document_id = str(job.get("document_id") or "")
    version = str(job.get("policy_version") or "")
    policy = controls.get_policy(document_id, version)
    if not policy:
        raise PolicyEngineError("Policy version was not found")
    stored_job = controls.get_policy_job(job_id)
    existing_controls = controls.policy_controls(document_id, version)
    if stored_job and stored_job.get("status") == "completed" and existing_controls:
        if stored_job.get("errors"):
            controls.update_policy_job(job_id, errors=[])
        return existing_controls
    controls.update_policy_job(job_id, status="running", phase="Extracting policy source")
    try:
        if policy.get("source_pending"):
            content, filename, media_type = fetch_public_policy(str(policy.get("source_url") or ""))
            policy = controls.replace_policy_source(document_id, version, content, filename)
            policy["media_type"] = media_type
        source = controls.get_blob(str(policy.get("source_blob") or ""))
        if not source:
            raise PolicyEngineError("Policy source artifact is unavailable")
        extraction = extract_document(str(policy.get("filename") or "source.txt"), source)
        controls.save_policy_extraction(document_id, version, text=extraction.text, clauses=extraction.clauses)
        controls.update_policy_job(job_id, status="running", phase="Generating and validating proposed controls")
        model = interpreter or AzureOpenAIPolicyInterpreter()
        proposal = assess_policy_proposal(model.interpret(policy, extraction.text, extraction.clauses), extraction.clauses)
        controls.save_policy_analysis(document_id, version, proposal)
        saved: list[dict[str, Any]] = []
        for raw_control in proposal["controls"]:
            if not isinstance(raw_control, dict):
                continue
            candidate = compile_proposal(raw_control, policy, extraction.clauses)
            # Retried deliveries may resume after controls were written but
            # before the message acknowledgement.  Reuse that exact proposal.
            previous = controls.get_control(candidate["control_id"], candidate["version"])
            if previous and previous.get("policy_document_id") == policy["document_id"] and previous.get("policy_version") == version:
                saved.append(previous)
                continue
            # Compatibility for controls generated before identifiers were
            # namespaced.  Titles are validated against this policy/version.
            legacy = next(
                (
                    item
                    for item in controls.policy_controls(document_id, version)
                    if item.get("title") == candidate["title"]
                    and str((item.get("source_reference") or {}).get("excerpt") or "")
                    == str((candidate.get("source_reference") or {}).get("excerpt") or "")
                ),
                None,
            )
            saved.append(legacy or controls.save_control(candidate))
        if not saved:
            raise PolicyEngineError("No policy controls were proposed")
        reconciled_policy = controls.reconcile_policy_coverage(document_id, version)
        needs_clarification = (
            not bool(reconciled_policy.get("coverage_complete"))
            or any(item["state"] == "needs_clarification" for item in saved)
        )
        controls.update_policy_state(
            document_id,
            version,
            status="needs_clarification" if needs_clarification else "ready",
            ingestion_status="completed",
        )
        controls.update_policy_job(job_id, status="completed", phase="Proposed controls are ready", control_count=len(saved), errors=[])
        controls.audit("policy.ingestion-completed", "policy-engine", {"job_id": job_id, "control_count": len(saved)})
        return saved
    except Exception as exc:
        current = controls.get_policy(document_id, version)
        if current:
            controls.update_policy_state(document_id, version, status="needs_clarification", ingestion_status="failed")
        controls.update_policy_job(job_id, status="failed", phase="Policy ingestion failed", errors=[str(exc)[:1000]])
        raise
